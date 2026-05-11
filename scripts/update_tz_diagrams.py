"""
Добавляет раздел с архитектурными диаграммами в TZ_REPO_module_v1_11052026.docx.
Запуск: python scripts/update_tz_diagrams.py
"""
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = Path(__file__).parent.parent
DOCX_PATH = BASE / "TZ_REPO_module_v1_11052026.docx"
BACKUP_PATH = BASE / "TZ_REPO_module_v1_11052026.bak.docx"

# Backup
shutil.copy2(DOCX_PATH, BACKUP_PATH)
print(f"Backup created: {BACKUP_PATH}")

doc = Document(DOCX_PATH)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def add_paragraph(doc, text, bold=False, italic=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def add_code_block(doc, text):
    """Add a monospace code-style paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    return p


def add_table_row(table, cells):
    row = table.add_row()
    for i, text in enumerate(cells):
        row.cells[i].text = text
        for para in row.cells[i].paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)
    return row


TABLE_STYLE = "Normal Table"


def set_header_row(row, bg_color="1F4E79"):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), bg_color)
        tcPr.append(shd)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9)


# ─── Add page break before new section ───────────────────────
doc.add_page_break()

# ─── Section heading ─────────────────────────────────────────
add_heading(doc, "15. Архитектурные диаграммы (C4 и UML)", level=1)

add_paragraph(doc,
    "Настоящий раздел содержит архитектурные диаграммы Модуля обработки сделок РЕПО, "
    "выполненные в нотации C4 Model и UML. Исходные файлы диаграмм в формате Mermaid (.mmd) "
    "поставляются вместе с исходным кодом проекта и могут быть отрендерены в любом "
    "Mermaid-совместимом инструменте (Mermaid Live Editor, GitLab, GitHub, VS Code).",
    size=10
)

# ─── Table of diagrams ────────────────────────────────────────
add_heading(doc, "15.1 Перечень диаграмм", level=2)

tbl = doc.add_table(rows=1, cols=4)
tbl.style = TABLE_STYLE
hdr = tbl.rows[0]
hdr.cells[0].text = "Файл"
hdr.cells[1].text = "Тип"
hdr.cells[2].text = "Нотация"
hdr.cells[3].text = "Содержание"
set_header_row(hdr)

diagrams = [
    ("d1_c4_context.mmd",   "C4 Level 1",    "C4 Context",   "Контекстная диаграмма: внешние системы и акторы"),
    ("d2_c4_container.mmd", "C4 Level 2",    "C4 Container", "Контейнеры: API, BatchProcessor, сервисы, БД"),
    ("d3_state_trade.mmd",  "UML State",     "stateDiagram", "Жизненный цикл сделки РЕПО (NEW→ACTIVE→CLOSED)"),
    ("d4_seq_batch.mmd",    "UML Sequence",  "sequenceDiagram","Ночная пакетная загрузка (чанки, SAVEPOINT)"),
    ("d5_er_model.mmd",     "ER-диаграмма",  "erDiagram",    "Модель данных: все таблицы и связи"),
    ("d6_flow_eod.mmd",     "Flowchart",     "flowchart TD", "EOD-расчёт: SOD → Leg1 → Leg2 → отчёт"),
    ("d7_seq_record.mmd",   "UML Sequence",  "sequenceDiagram","Обработка одной записи внутри SAVEPOINT"),
    ("d8_c4_component.mmd", "C4 Level 3",    "C4 Component", "Компоненты модуля и их зависимости"),
]
for row_data in diagrams:
    add_table_row(tbl, row_data)

# ─── D1: Context ─────────────────────────────────────────────
add_heading(doc, "15.2 D1 — Контекстная диаграмма (C4 Level 1)", level=2)
add_paragraph(doc,
    "Показывает Модуль РЕПО в контексте взаимодействующих внешних систем: "
    "Торговой системы (источник сделок), Расчётной системы НКЦ (SOD-остатки, подтверждения Leg 2), "
    "Депозитария (справочники), Системы мониторинга и Оператора.",
    size=10
)
add_paragraph(doc, "Файл: d1_c4_context.mmd", italic=True, size=9)

# ─── D2: Container ───────────────────────────────────────────
add_heading(doc, "15.3 D2 — Диаграмма контейнеров (C4 Level 2)", level=2)
add_paragraph(doc,
    "Раскрывает внутреннюю структуру Модуля РЕПО. Ключевые контейнеры:",
    size=10
)
containers = [
    ("REST API (FastAPI + Uvicorn)",
     "GET /trades, /positions, /obligations, /reports; DELETE /trades/{id}; POST /batch/load"),
    ("BatchProcessor (batch/processor.py)",
     "Чанковая обработка JSON Lines. SAVEPOINT на запись (бизнес-ошибки). "
     "ROLLBACK чанка (системные ошибки). Повтор до 3 раз."),
    ("TradeService (services/trade_service.py)",
     "Маппинг ролей party_1/party_2. Fallback DEFAULT_SELLER. "
     "Валидация справочников. Расчёт leg2_amount (HALF_UP)."),
    ("PositionService (services/position_service.py)",
     "Pre-check balance >= 0 (1-й рубеж). Применение дельт Leg 1/Leg 2. "
     "Future Obligations. Аудит-лог (SHA-256 hash-chain)."),
    ("SQLite (демо-режим, aiosqlite)",
     "Полная эмуляция без внешних зависимостей. "
     "CHECK balance>=0 через BEFORE INSERT/UPDATE триггеры. MODE=demo."),
    ("PostgreSQL 15+ (прод-режим, asyncpg)",
     "Партиционирование raw_trades. Нативный CHECK constraint. JSONB. MODE=production."),
]
tbl2 = doc.add_table(rows=1, cols=2)
tbl2.style = TABLE_STYLE
hdr2 = tbl2.rows[0]
hdr2.cells[0].text = "Контейнер"
hdr2.cells[1].text = "Назначение"
set_header_row(hdr2)
for name, desc in containers:
    add_table_row(tbl2, [name, desc])

add_paragraph(doc, "Файл: d2_c4_container.mmd", italic=True, size=9)

# ─── D3: State ───────────────────────────────────────────────
add_heading(doc, "15.4 D3 — Жизненный цикл сделки (UML State)", level=2)
add_paragraph(doc,
    "Диаграмма состояний сделки РЕПО. В batch-режиме (MVP) сделка сразу переходит "
    "в статус ACTIVE, минуя NEW, поскольку trade_date = leg1_settlement_date = EOD-дата.",
    size=10
)
states = [
    ("NEW → ACTIVE",    "trade_date = EOD-дата, pre-check OK, проводки Leg 1 применены"),
    ("NEW → CANCELLED", "DELETE /trades/{id}, current_date < leg1_settlement_date"),
    ("NEW → REJECTED",  "Pre-check FAIL (INSUFFICIENT_BALANCE) или VALIDATION_ERROR"),
    ("ACTIVE → CLOSED", "LEG2_SETTLED получен, проводки Leg 2 применены"),
    ("ACTIVE → DEFAULTED", "[Этап 2] LEG2_FAILED — ручной разбор в MVP"),
    ("ACTIVE → EARLY_TERMINATED", "[Этап 2] Досрочное закрытие"),
]
tbl3 = doc.add_table(rows=1, cols=2)
tbl3.style = TABLE_STYLE
hdr3 = tbl3.rows[0]
hdr3.cells[0].text = "Переход"
hdr3.cells[1].text = "Условие / действие"
set_header_row(hdr3)
for t, c in states:
    add_table_row(tbl3, [t, c])
add_paragraph(doc, "Файл: d3_state_trade.mmd", italic=True, size=9)

# ─── D4: Batch Sequence ──────────────────────────────────────
add_heading(doc, "15.5 D4 — Ночная пакетная загрузка (UML Sequence)", level=2)
add_paragraph(doc,
    "Детальная последовательность ночной пакетной загрузки. "
    "Участники: BatchProcessor, TradeService, PositionService, БД, Мониторинг.",
    size=10
)
add_paragraph(doc, "Алгоритм обработки чанка:", bold=True, size=10)
steps = [
    "1. BEGIN TRANSACTION (чанк)",
    "2. Для каждой записи: SAVEPOINT sp_record_{chunk}_{n}",
    "3. Парсинг JSON → IncomingTrade (Pydantic-валидация)",
    "4. Dedup-проверка по idempotency_key (external_trade_id + trade_date)",
    "5. validate_trade(): проверка Participant и Instrument в справочниках",
    "6. map_roles(): SECURITY_SELLER / SECURITY_BUYER / DEFAULT_SELLER (fallback + WARNING)",
    "7. calculate_leg2_amount(): ROUND(sum*rate*days/365, 2, HALF_UP)",
    "8. Pre-check: SELECT balance FOR UPDATE → проверка balance + delta >= 0",
    "9. INSERT raw_trades, INSERT trades, UPDATE positions, INSERT audit_log, INSERT obligations",
    "10. RELEASE SAVEPOINT → committed++",
    "Бизнес-ошибка (шаги 3–8): ROLLBACK TO SAVEPOINT → rejected_trades → CONTINUE",
    "Системная ошибка: ROLLBACK чанка → повтор до max_retries=3 → failed_chunks++",
    "11. COMMIT чанка",
]
for s in steps:
    p = doc.add_paragraph(s, style="List Paragraph")
    for run in p.runs:
        run.font.size = Pt(9)
add_paragraph(doc, "Файл: d4_seq_batch.mmd", italic=True, size=9)

# ─── D5: ER Model ────────────────────────────────────────────
add_heading(doc, "15.6 D5 — ER-модель данных", level=2)
add_paragraph(doc,
    "Реляционная модель данных Модуля РЕПО. "
    "Типы данных адаптированы для двух режимов: "
    "PostgreSQL (UUID, JSONB, BIGSERIAL, DECIMAL) и SQLite (TEXT, JSON, INTEGER, REAL).",
    size=10
)
tables_desc = [
    ("participants",        "Справочник участников клиринга"),
    ("instruments",         "Справочник инструментов (ISIN, repo_eligible, day_count_convention)"),
    ("raw_trades",          "Неизменяемый журнал входящих записей (append-only, dedup-индекс)"),
    ("trades",              "Обработанные сделки РЕПО (все поля, статус, leg2_amount)"),
    ("positions",           "EOD/SOD позиции участников (CHECK balance >= 0)"),
    ("future_obligations",  "Реестр будущих обязательств по Leg 2 (4 записи на сделку)"),
    ("position_audit_log",  "Аудит-лог изменений позиций (append-only, SHA-256 hash-chain)"),
    ("rejected_trades",     "Отклонённые записи с типом и деталями ошибки"),
    ("load_reports",        "Отчёты о пакетных загрузках (UNIQUE по eod_date)"),
]
tbl5 = doc.add_table(rows=1, cols=2)
tbl5.style = TABLE_STYLE
hdr5 = tbl5.rows[0]
hdr5.cells[0].text = "Таблица"
hdr5.cells[1].text = "Назначение"
set_header_row(hdr5)
for name, desc in tables_desc:
    add_table_row(tbl5, [name, desc])
add_paragraph(doc, "Файл: d5_er_model.mmd", italic=True, size=9)

# ─── D6: EOD Flow ────────────────────────────────────────────
add_heading(doc, "15.7 D6 — EOD-расчёт позиций (Flowchart)", level=2)
add_paragraph(doc,
    "Блок-схема алгоритма расчёта позиций на конец торгового дня (EOD). "
    "Реализована в scripts/run_demo.py и batch/processor.py.",
    size=10
)
eod_steps = [
    "load_sod_balances() — загрузка SOD-остатков в positions (status=SOD)",
    "copy_sod_to_eod() — копирование SOD → EOD как стартовая точка дня",
    "run_batch_load() — пакетная загрузка сделок, применение Leg 1 в процессе",
    "Применение Leg 2 для leg2_settlement_date = D (при наличии LEG2_SETTLED)",
    "Проверка inconsistent-позиций → алерт при наличии",
    "Формирование EOD Position Report (CSV/JSON), Future Obligations Register, Load Report",
]
for s in eod_steps:
    p = doc.add_paragraph(s, style="List Paragraph")
    for run in p.runs:
        run.font.size = Pt(9)
add_paragraph(doc, "Файл: d6_flow_eod.mmd", italic=True, size=9)

# ─── D7: Record Sequence ─────────────────────────────────────
add_heading(doc, "15.8 D7 — Обработка одной записи (UML Sequence)", level=2)
add_paragraph(doc,
    "Детальная последовательность обработки одной JSON-строки внутри активного SAVEPOINT. "
    "Участники: BatchProcessor._process_record(), TradeService, PositionService, БД.",
    size=10
)
add_paragraph(doc,
    "Последовательность: json.loads → IncomingTrade (Pydantic) → dedup → validate_trade → "
    "map_roles → calculate_leg2_amount → pre-check x4 (participant/counterparty × securities/cash) → "
    "INSERT trades → UPDATE positions → INSERT audit_log (SHA-256) → INSERT obligations x4 → "
    "RELEASE SAVEPOINT.",
    size=10
)
add_paragraph(doc, "Файл: d7_seq_record.mmd", italic=True, size=9)

# ─── D8: Component ───────────────────────────────────────────
add_heading(doc, "15.9 D8 — Диаграмма компонентов (C4 Level 3)", level=2)
add_paragraph(doc,
    "Детализирует внутреннюю структуру Модуля РЕПО до уровня отдельных Python-модулей "
    "и их зависимостей.",
    size=10
)
components = [
    ("BatchProcessor",    "batch/processor.py",              "Чанки, SAVEPOINT, retry, счётчики"),
    ("TradeService",      "services/trade_service.py",       "map_roles, validate_trade, build_trade_create, insert_trade, cancel_trade"),
    ("PositionService",   "services/position_service.py",    "apply_position_delta, apply_leg1_settlements, create_future_obligations, load_sod_balances"),
    ("calc.py",           "utils/calc.py",                   "calculate_leg2_amount (HALF_UP, ACT/365, ACT/360)"),
    ("hashing.py",        "utils/hashing.py",                "compute_chain_hash (SHA-256 hash-chain)"),
    ("ORM Models",        "db/orm.py",                       "PortableUUID, PortableJSON, SQLite-триггеры"),
    ("DB Base",           "db/base.py",                      "get_engine, get_session, init_db, register_sqlite_events"),
    ("REST API Routes",   "api/routes.py",                   "FastAPI эндпоинты GET/DELETE/POST"),
    ("Config",            "config.py + config.yaml + .env",  "MODE=demo/production, get_db_url, get_chunk_size"),
]
tbl8 = doc.add_table(rows=1, cols=3)
tbl8.style = TABLE_STYLE
hdr8 = tbl8.rows[0]
hdr8.cells[0].text = "Компонент"
hdr8.cells[1].text = "Файл"
hdr8.cells[2].text = "Ответственность"
set_header_row(hdr8)
for name, f, desc in components:
    add_table_row(tbl8, [name, f, desc])
add_paragraph(doc, "Файл: d8_c4_component.mmd", italic=True, size=9)

# ─── Save ─────────────────────────────────────────────────────
doc.save(DOCX_PATH)
print(f"Updated: {DOCX_PATH}")
print("Section '15. Архитектурные диаграммы' added successfully.")
